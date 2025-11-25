"""
Document Parser Service for FileGPT
Supports text, code, PDF, and DOCX files with intelligent binary filtering.
"""

import os
from pathlib import Path
from typing import Optional

# Lazy imports for document libraries
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Supported file extensions
TEXT_EXTENSIONS = {
    '.txt', '.md', '.markdown', '.rst',
    '.json', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf',
    '.html', '.htm', '.css', '.scss', '.sass', '.less',
    '.log', '.csv', '.tsv'
}

CODE_EXTENSIONS = {
    '.py', '.pyw', '.pyi',  # Python
    '.js', '.jsx', '.ts', '.tsx', '.mjs', '.cjs',  # JavaScript/TypeScript
    '.java', '.kt', '.kts',  # JVM languages
    '.c', '.h', '.cpp', '.hpp', '.cc', '.cxx',  # C/C++
    '.rs',  # Rust
    '.go',  # Go
    '.rb',  # Ruby
    '.php',  # PHP
    '.swift',  # Swift
    '.sh', '.bash', '.zsh', '.fish',  # Shell
    '.ps1', '.psm1',  # PowerShell
    '.r', '.R',  # R
    '.sql',  # SQL
    '.m', '.mm',  # Objective-C
    '.cs',  # C#
    '.vb',  # Visual Basic
    '.lua',  # Lua
    '.pl', '.pm',  # Perl
    '.scala',  # Scala
    '.clj', '.cljs',  # Clojure
    '.ex', '.exs',  # Elixir
    '.erl', '.hrl',  # Erlang
    '.dart',  # Dart
    '.groovy',  # Groovy
    '.vim',  # Vim script
}

DOCUMENT_EXTENSIONS = {
    '.pdf',
    '.docx', '.doc'
}

# Combined set of all supported extensions
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | CODE_EXTENSIONS | DOCUMENT_EXTENSIONS


def get_file_content(file_path: str) -> Optional[str]:
    """
    Extract text content from a file.
    
    Args:
        file_path: Absolute path to the file
        
    Returns:
        Extracted text content, or None if the file is unsupported/binary
    """
    if not os.path.exists(file_path):
        return None
    
    # Skip if it's a directory
    if os.path.isdir(file_path):
        return None
    
    # Get file extension
    extension = Path(file_path).suffix.lower()
    
    # Check if extension is supported
    if extension not in SUPPORTED_EXTENSIONS:
        return None
    
    try:
        # Handle text and code files
        if extension in TEXT_EXTENSIONS or extension in CODE_EXTENSIONS:
            return _read_text_file(file_path)
        
        # Handle PDF files
        elif extension == '.pdf':
            return _read_pdf(file_path)
        
        # Handle DOCX files
        elif extension in {'.docx', '.doc'}:
            return _read_docx(file_path)
        
        return None
        
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return None


def _read_text_file(file_path: str) -> Optional[str]:
    """Read plain text or code file with UTF-8 encoding."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            
        # Filter out files that are likely binary despite having text extension
        # Check for high ratio of non-printable characters
        if content:
            non_printable = sum(1 for c in content[:1000] if ord(c) < 32 and c not in '\n\r\t')
            if non_printable > 100:  # More than 10% non-printable in first 1000 chars
                return None
                
        return content if content.strip() else None
        
    except UnicodeDecodeError:
        return None
    except Exception as e:
        print(f"Error reading text file {file_path}: {e}")
        return None


def _read_pdf(file_path: str) -> Optional[str]:
    """Extract text from PDF file using pypdf."""
    if not PYPDF_AVAILABLE:
        print("pypdf not installed. Cannot read PDF files.")
        return None
    
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        content = '\n'.join(text_parts)
        return content if content.strip() else None
        
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return None


def _read_docx(file_path: str) -> Optional[str]:
    """Extract text from DOCX file using python-docx."""
    if not DOCX_AVAILABLE:
        print("python-docx not installed. Cannot read DOCX files.")
        return None
    
    try:
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        content = '\n'.join(text_parts)
        return content if content.strip() else None
        
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return None


def is_supported_file(file_path: str) -> bool:
    """
    Check if a file is supported for indexing.
    
    Args:
        file_path: Path to the file
        
    Returns:
        True if the file extension is supported
    """
    extension = Path(file_path).suffix.lower()
    return extension in SUPPORTED_EXTENSIONS
